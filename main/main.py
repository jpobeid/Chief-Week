import datetime
import global_vars as globals
import functions as funcs
import didactics as lecs
import dept_qg as qg
import docx
from docx.shared import RGBColor
import numpy as np
import pandas as pd

if not funcs.screen_assets():
    error_message = f'Requires {globals.FILENAME_SETTINGS} and {globals.FILENAME_SCHEDULE} in /{globals.PATH_ASSETS}'
    raise FileNotFoundError(error_message)

document = docx.Document()
settings = funcs.make_settings_dict()
settings_keys = list(settings.keys())

#---Start document---
#---Intro/News---
document.add_heading('News', 1)

document.add_paragraph('***')

#---Reminders---
document.add_heading('Reminders', 1)

i_reminders = np.argwhere([e.find('reminder') != -1 for e in settings_keys]).flatten()
p = document.add_paragraph()
for j, i in enumerate(i_reminders):
    prefix = '\n' if j != 0 else ''
    p.add_run(prefix + str(j + 1) + '. ' + settings.get(settings_keys[i]))

#---Lecture Zoom Info---
document.add_heading('Lectures Zoom Info', 1)

p = document.add_paragraph()
p.add_run(settings.get('host_link') + '\n')
p.add_run('Password: ' + settings.get('host_pass') + '\n')
p.add_run('Meeting ID: ' + settings.get('host_id') + '\n')
p.add_run('Dial-in: ' + settings.get('host_dial'))

#---Schedule---
document.add_heading('Schedule', 1)

df_lec = lecs.get_weekly_lectures()
next_week = funcs.get_next_week()
for i, e in enumerate(globals.SCHEDULE_DAYS):
    date = str(next_week[i].month) + '/' + str(next_week[i].day)
    p = document.add_paragraph()
    r = p.add_run(f'{e} ({date}):')
    r.bold = True
    r.underline = True
    document.add_paragraph('***')
    if next_week[i] in list(map(lambda x: x.date(), df_lec.index)):
        index_date = list(map(lambda x: x.date(), df_lec.index)).index(next_week[i])
        lectures = df_lec.loc[df_lec.index[index_date]]
        p = document.add_paragraph()
        if type(lectures) == pd.Series:
            lecture = lectures
            funcs.add_lecture_runs(p, lecture, new_line=False)
        else:
            for j in range(lectures.shape[0]):
                to_new_line = True if j < lectures.shape[0] - 1 else False
                lecture = lectures.iloc[j]
                funcs.add_lecture_runs(p, lecture, new_line=to_new_line)
    if e == settings.get('seminar_weekday'):
        time_seminar = settings.get('seminar_time')[:-2] + ' ' + settings.get('seminar_time')[-2:]
        date_seminar = next_week[i]
        url_seminar = settings.get('seminar_calendar').replace('mm-dd-yyyy', next_week[0].strftime('%m-%d-%Y'))
        seminar_text = funcs.get_seminar_text(url_seminar, date_seminar, time_seminar)
        p = document.add_paragraph(f'8:00-9:00 - Cancer series seminar:\n{seminar_text}')
        seminar_link = settings.get('seminar_link')
        seminar_pass = settings.get('seminar_pass')
        p.add_run(f'\nGeneral link: {seminar_link}')
        p.add_run(f'\nPassword: {seminar_pass}')
    elif e == 'Thursday':
        p = document.add_paragraph('7:30-8:30 - Chart Rounds: (Please see dosi email, links are reference only)')
        link_875 = settings.get('875_link')
        pass_875 = settings.get('875_pass')
        link_900 = settings.get('900_link')
        pass_900 = settings.get('900_pass')
        p.add_run(f'\n875-BW link: {link_875}\n')
        p.add_run(f'875-BW password: {pass_875}\n')
        p.add_run(f'900-BW link: {link_900}\n')
        p.add_run(f'900-BW password: {pass_900}')

#---Groups---
document.add_heading('Groups', 1)

# Request updated sheet from Qgenda
html = qg.get_html_qg(settings.get('qgenda'), globals.HTTP_HEADERS)

df_res = qg.make_df_qg(html, globals.COLOR_RESIDENT, True).T
df_attg = qg.make_df_qg(html, globals.COLOR_RESIDENT, False).T
off_campus = qg.get_dict_var(df_res, 'Off-Campus', next_week)
first_call = qg.get_dict_var(df_res, '1st Call', next_week)
second_call = qg.get_dict_var(df_attg, '2nd Call', next_week)

# Call
p_i = document.add_paragraph()
funcs.add_group_run(p_i, 'Call:')
p_i.add_run('1st call:\n').bold = True
for i, e in enumerate(funcs.format_group_dict(first_call, True)):
    p_i.add_run(funcs.clean_name(e[0]) + ': ' + str(e[1]).replace("'", '') + funcs.get_suffix(i, first_call))
p_i.add_run('\n')
p_i.add_run('2nd call:\n').bold = True
for i, e in enumerate(funcs.format_group_dict(second_call, True)):
    p_i.add_run(funcs.clean_name(e[0]) + ': ' + str(e[1]).replace("'", '') + funcs.get_suffix(i, second_call))

# Off-campus
p_i = document.add_paragraph()
funcs.add_group_run(p_i, 'Off-campus:')
for i, e in enumerate(funcs.format_group_dict(off_campus, False)):
    p_i.add_run(funcs.clean_name(e[0]) + ': ' + str(e[1]).replace("'", '') + funcs.get_suffix(i, off_campus))

# Qgenda reference
p_i = document.add_paragraph()
funcs.add_group_run(p_i, 'Qgenda reference:')
p_i.add_run(settings.get('qgenda'))

# Varian scholars
p_i = document.add_paragraph()
funcs.add_group_run(p_i, 'Varian scholars: (Please submit 2 business days prior to end of month)')
groups = settings.get('varian_groups').split('/')
n_groups = len(groups)
order = np.array(settings.get('varian_order').split('/')).astype(int)
if order[0] != 1 or not np.all(np.diff(order) == np.repeat(1, n_groups - 1)):
    raise ValueError('Order settings is not appropriately set, must start at 1 and increment to match group length')
i_month = datetime.date.today().month - 1
for i in range(n_groups):
    symbol = globals.MONTH_SYMBOLS[(i_month + i) % 12]
    group_i = (i_month + i) % n_groups
    group = groups[int(order[group_i]) - 1].split(',')
    group[0] = group[0] + ' (Chief)'
    r = p_i.add_run(symbol + ': ' + ', '.join(group) + funcs.get_suffix(i, groups))
    if i == 0:
        font = r.font
        font.color.rgb = RGBColor(0xff, 0x00, 0x00)

#---Sign document---
document.add_paragraph('Templated with \u2665 by Python')

#---Save document---
document.save(globals.FILENAME_DOCX)

#---Manually close terminal---
input('Press any key to exit...')